"""
report_generator.py — L2 级测试报告 Word 文档生成器

按公司内部 L2 测试报告模板生成 .docx 文件。
使用 python-docx 库操作 Word 文档。

模板结构:
- 页眉: 公司名称 + 内部L2
- 标题: 项目测试报告
- 文件信息表
- 评审/批准人员
- 版次历史表
- 测试需求来源
- 测试策略
- 测试大项执行表
- 测试结果 (遗留问题 / 专项测试)
- 测试结论
"""

import io
import os
from datetime import datetime
from typing import Dict, List, Optional
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from api.models.test_report_db import get_report


def _set_cell_shading(cell, color: str):
    """设置单元格背景色"""
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def _set_cell_border(cell, **kwargs):
    """设置单元格边框"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge, val in kwargs.items():
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), val.get("val", "single"))
        element.set(qn("w:sz"), val.get("sz", "4"))
        element.set(qn("w:color"), val.get("color", "000000"))
        tcBorders.append(element)
    tcPr.append(tcBorders)


def _add_formatted_table(doc, headers: List[str], rows: List[List[str]], col_widths: List[float] = None):
    """添加带格式的表格（表头灰底）"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 表头
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
        _set_cell_shading(cell, "D9E2F3")

    # 数据行
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.rows[r + 1].cells[c]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)

    return table


def generate_l2_report(report_id: int, output_path: str = None) -> str:
    """
    生成 L2 测试报告 .docx 文件。

    参数:
        report_id: 数据库中的报告 ID
        output_path: 输出路径，不传则返回 bytes (用于 HTTP 下载)

    返回: 文件路径
    """
    data = get_report(report_id)
    if not data:
        raise ValueError(f"报告 {report_id} 不存在")

    doc = Document()

    # ---- 页面设置 ----
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

    # ---- 页眉 ----
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = hp.add_run("深圳市睿联技术有限公司                                         内部L2")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(128, 128, 128)

    # ---- 标题 ----
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(f"\n{data.get('project_name', '项目')} {data.get('firmware_version', '')} 正式迭代版本测试报告")
    run.bold = True
    run.font.size = Pt(16)

    doc.add_paragraph()  # 空行

    # ---- 文件信息 ----
    doc.add_heading("一、文件信息", level=2)
    _add_formatted_table(doc,
        ["项目", "内容"],
        [
            ["文件编号", data.get("file_number", "")],
            ["版次", data.get("version", "")],
            ["生效日期", data.get("effective_date", "")],
            ["归口部门", "测试部"],
            ["编制人", data.get("preparer", "")],
            ["测试状态", data.get("test_status", "")],
        ],
    )

    doc.add_paragraph()

    # ---- 评审/批准人员 ----
    doc.add_heading("二、评审/批准人员", level=2)
    _add_formatted_table(doc,
        ["角色", "人员"],
        [
            ["产品经理", data.get("product_manager", "")],
            ["项目经理", data.get("project_manager", "")],
            ["开发负责人", data.get("dev_lead", "")],
            ["硬件负责人", data.get("hardware_lead", "")],
            ["结构负责人", data.get("structure_lead", "")],
            ["测试负责人", data.get("test_lead", "")],
        ],
    )

    doc.add_paragraph()

    # ---- 发布范围 ----
    doc.add_heading("三、发布范围", level=2)
    p = doc.add_paragraph("公司内部 （深圳研发中心 + 成都研发中心）")

    doc.add_paragraph()

    # ---- 版次历史 ----
    doc.add_heading("四、版次历史", level=2)
    _add_formatted_table(doc,
        ["版本", "日期", "修订描述", "修订人"],
        [[data.get("version", ""), data.get("effective_date", ""), "初始版本", data.get("preparer", "")]],
    )

    doc.add_paragraph()

    # ---- 测试需求来源 ----
    doc.add_heading("五、测试需求来源", level=2)
    _add_formatted_table(doc,
        ["项目", "内容"],
        [
            ["项目背景", f"{data.get('project_name', '')} 正式迭代版本测试"],
            ["产品型号", data.get("product_model", "")],
            ["SoC", data.get("soc_info", "")],
            ["Switch", data.get("switch_info", "")],
            ["Flash", data.get("flash_info", "")],
            ["DDR", data.get("ddr_info", "")],
            ["固件版本", data.get("firmware_version", "")],
            ["测试报告类型", "内部L2"],
            ["测试范围", data.get("test_range", "")],
        ],
    )

    doc.add_paragraph()

    # ---- 测试策略 ----
    doc.add_heading("六、测试策略", level=2)
    strategy = data.get("test_strategy", "") or (
        "1. 功能测试：逐项验证各功能模块是否按需求正常工作。\n"
        "2. 专项测试：图像、音频、Wi-Fi、拷机等专项深度测试。\n"
        "3. 兼容性测试：验证各平台、各浏览器的兼容性。\n"
        "4. 升级测试：验证固件升级路径的稳定性。"
    )
    for line in strategy.split("\n"):
        if line.strip():
            doc.add_paragraph(line.strip())

    doc.add_paragraph()

    # ---- 测试大项和执行情况 ----
    doc.add_heading("七、测试大项和执行情况", level=2)
    test_items = data.get("test_items", [])
    if test_items:
        _add_formatted_table(doc,
            ["测试大项", "执行状态", "执行结果"],
            [[item.get("category", ""), item.get("status", ""), item.get("result", "")]
             for item in test_items],
        )

    doc.add_paragraph()

    # ---- 测试结果 ----
    doc.add_heading("八、测试结果", level=2)

    # 主要遗留问题 (Severe/Critical)
    bugs = data.get("bugs", [])
    severe_bugs = [b for b in bugs if b.get("severity", "") in ("Critical", "Severe")]
    doc.add_heading("8.1 主要遗留问题 (Severe / Critical)", level=3)
    if severe_bugs:
        _add_formatted_table(doc,
            ["Bug编号", "Bug标题", "严重程度", "状态", "解决方案"],
            [[b.get("bug_number", ""), b.get("bug_title", ""),
              b.get("severity", ""), b.get("bug_status", ""), b.get("solution", "")]
             for b in severe_bugs],
        )
    else:
        doc.add_paragraph("无 Critical / Severe 级别遗留问题。")

    doc.add_paragraph()

    # 全部遗留问题清单
    doc.add_heading("8.2 全部遗留问题清单", level=3)
    if bugs:
        _add_formatted_table(doc,
            ["Bug编号", "Bug标题", "严重程度", "Bug状态", "解决方案"],
            [[b.get("bug_number", ""), b.get("bug_title", ""),
              b.get("severity", ""), b.get("bug_status", ""), b.get("solution", "")]
             for b in bugs],
        )

    doc.add_paragraph()

    # 各专项测试结果
    sub_items = data.get("sub_items", [])
    doc.add_heading("8.3 各专项测试结果", level=3)
    if sub_items:
        _add_formatted_table(doc,
            ["专项名称", "测试结果", "详细描述"],
            [[item.get("sub_name", ""), item.get("result", ""), item.get("detail", "")]
             for item in sub_items],
        )

    doc.add_paragraph()

    # ---- 测试结论 ----
    doc.add_heading("九、测试结论", level=2)
    summary = data.get("summary", "") or (
        f"本次测试针对 {data.get('project_name', '项目')} {data.get('firmware_version', '')} "
        f"版本进行了全面验证。测试范围覆盖 {data.get('test_range', '各功能模块')}。\n\n"
        f"遗留 Bug 共 {len(bugs)} 个，其中 Critical/Severe 级别 {len(severe_bugs)} 个。"
    )
    doc.add_paragraph(summary)

    # ---- 保存 ----
    if not output_path:
        output_dir = os.path.join(os.path.dirname(__file__), "..", "..", "reports")
        os.makedirs(output_dir, exist_ok=True)
        output_path = os.path.join(output_dir, f"L2_Report_{data.get('file_number', str(report_id))}.docx")

    doc.save(output_path)
    return output_path


def generate_report_bytes(report_id: int) -> bytes:
    """生成报告并返回 bytes（用于 HTTP 流式下载）"""
    data = get_report(report_id)
    if not data:
        raise ValueError(f"报告 {report_id} 不存在")

    tmp_path = os.path.join(
        os.path.dirname(__file__), "..", "..", "reports",
        f"_tmp_{report_id}.docx"
    )
    generate_l2_report(report_id, tmp_path)
    with open(tmp_path, "rb") as f:
        content = f.read()
    os.remove(tmp_path)
    return content
